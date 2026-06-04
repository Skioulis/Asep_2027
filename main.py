import json
import os
import re
from docx import Document

BLUE_COLORS = {'002060', '001F5F'}
RED_COLORS = {'FF0000', 'EE0000', 'C00000', 'DA0000', 'ED0000', 'CC0000', 'B80000', '9F0000'}
ANSWER_PREFIX = re.compile(r'^[αβγδabcdABCD][.)]\s*')


def _run_is_red(run):
    return bool(run.font.color.rgb and str(run.font.color.rgb) in RED_COLORS)


def _para_is_blue(para):
    return any(
        run.font.color.rgb and str(run.font.color.rgb) in BLUE_COLORS
        for run in para.runs
    )


def _expand_paragraph(para):
    """
    Return a list of (text, is_red) items.
    A single paragraph may contain multiple newline-separated answer options
    when the document author merged them into one paragraph.
    """
    current_text = ""
    current_red = False
    segments = []

    for run in para.runs:
        run_red = _run_is_red(run)
        parts = run.text.split('\n')
        for i, part in enumerate(parts):
            current_text += part
            if run_red and part.strip():
                current_red = True
            if i < len(parts) - 1:
                if current_text.strip():
                    segments.append((current_text.strip(), current_red))
                current_text = ""
                current_red = False

    if current_text.strip():
        segments.append((current_text.strip(), current_red))

    return segments


def _clean_answer(text):
    return ANSWER_PREFIX.sub('', text).strip()


def extract_from_file(filepath):
    doc = Document(filepath)

    # Extract category name from document header
    category = os.path.splitext(os.path.basename(filepath))[0]
    for para in doc.paragraphs:
        if para.text.strip().startswith('Γνωστικό Αντικείμενο:'):
            category = para.text.strip().removeprefix('Γνωστικό Αντικείμενο:').strip()
            break

    # Collect flat list of (text, is_red) after the header
    items = []
    header_done = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if not header_done:
            if _para_is_blue(para):
                if text.startswith('Γνωστικό Αντικείμενο:'):
                    header_done = True
                continue
            continue
        if _para_is_blue(para):
            continue
        items.extend(_expand_paragraph(para))

    # Group every 5 items into a question block
    questions = []
    i = 0
    while i + 4 < len(items):
        q_text, q_red = items[i]
        answers_raw = items[i + 1: i + 5]

        # The question text itself should not be red
        if q_red:
            i += 1
            continue

        red_indices = [j for j, (_, red) in enumerate(answers_raw) if red]
        if len(red_indices) == 0:
            i += 1
            continue

        correct_idx = red_indices[0]
        answers = [_clean_answer(a[0]) for a in answers_raw]

        questions.append({
            "question": q_text,
            "answers": [
                {"option": chr(ord('a') + k), "text": answers[k]}
                for k in range(4)
            ],
            "correctAnswer": chr(ord('a') + correct_idx),
            "category": category,
        })
        i += 5

    return questions


def main():
    questions_dir = "questions"
    output_path = os.path.join("data", "questions.json")
    os.makedirs("data", exist_ok=True)

    all_questions = []
    q_id = 1

    for filename in sorted(os.listdir(questions_dir)):
        if not filename.endswith(".docx"):
            continue
        # Skip the announcement file
        if filename.startswith("26."):
            continue

        filepath = os.path.join(questions_dir, filename)
        print(f"Processing: {filename}")
        extracted = extract_from_file(filepath)
        for q in extracted:
            q["id"] = str(q_id)
            q_id += 1
        all_questions.extend(extracted)
        print(f"  → {len(extracted)} questions")

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(all_questions, f, ensure_ascii=False, indent=2)

    print(f"\nTotal: {len(all_questions)} questions saved to {output_path}")


if __name__ == "__main__":
    main()
